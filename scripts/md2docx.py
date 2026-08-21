"""Minimal Markdown -> .docx converter (no pandoc on this box).

Lives in scripts/ rather than a scratchpad because the scratchpad has been
cleared three times mid-session, and regenerating the TUBITAK report needs it.

Handles ATX headings, pipe tables, fenced code, blockquotes, bullet/numbered
lists, horizontal rules, and inline **bold** / *italic* / `code`. Every branch
of the main loop must advance `i` -- an earlier version did not on the heading
branch and hung.

Usage: python3 scripts/md2docx.py in.md out.docx
"""
import re
import sys

from docx import Document
from docx.shared import Inches, Pt

INLINE = re.compile(r"(\*\*.+?\*\*|\*[^*]+?\*|`[^`]+?`)", re.S)
IMAGE = re.compile(r"^!\[([^\]]*)\]\(([^)]+)\)\s*$")


def add_runs(par, text):
    """Adds `text` to `par`, honouring bold/italic/code spans."""
    for tok in INLINE.split(text):
        if not tok:
            continue
        if tok.startswith("**") and tok.endswith("**"):
            par.add_run(tok[2:-2]).bold = True
        elif tok.startswith("*") and tok.endswith("*"):
            par.add_run(tok[1:-1]).italic = True
        elif tok.startswith("`") and tok.endswith("`"):
            r = par.add_run(tok[1:-1])
            r.font.name = "Consolas"
            r.font.size = Pt(9)
        else:
            par.add_run(tok)


def gather(lines, i, n):
    """Consumes wrapped continuation lines following a list item.

    A list item's text can wrap just like a paragraph's, and a bold span may
    open on the item line and close on a continuation line. Returns the joined
    text and the new index.
    """
    buf = []
    while i < n:
        cur = lines[i].strip()
        if (not cur or cur.startswith(("```", "|", ">", "#", "!["))
                or re.match(r"^[-*+]\s+|^\d+\.\s+", cur)
                or (set(cur) <= {"-", "*", "_"} and len(cur) >= 3)):
            break
        buf.append(cur)
        i += 1
    return " ".join(buf), i


def split_row(line):
    """Splits one Markdown table row into stripped cell strings."""
    return [c.strip() for c in line.strip().strip("|").split("|")]


def main(src, dst):
    """Converts `src` Markdown to `dst` .docx."""
    lines = open(src, encoding="utf-8").read().splitlines()
    doc = Document()
    i, n = 0, len(lines)
    while i < n:
        s = lines[i].strip()

        if not s:
            i += 1
            continue

        if s.startswith("```"):
            i += 1
            buf = []
            while i < n and not lines[i].strip().startswith("```"):
                buf.append(lines[i])
                i += 1
            i += 1
            r = doc.add_paragraph().add_run("\n".join(buf))
            r.font.name = "Consolas"
            r.font.size = Pt(9)
            continue

        if set(s) <= {"-", "*", "_"} and len(s) >= 3:
            doc.add_paragraph("_" * 40)
            i += 1
            continue

        m = IMAGE.match(s)                                        # figure
        if m:
            try:
                doc.add_picture(m.group(2), width=Inches(6.2))
            except Exception as e:
                doc.add_paragraph(f"[figure missing: {m.group(2)} -- {e}]")
            i += 1
            continue

        m = re.match(r"^(#{1,6})\s+(.*)", s)
        if m:
            doc.add_heading(m.group(2), level=min(len(m.group(1)), 4))
            i += 1
            continue

        if s.startswith("|") and i + 1 < n and re.match(r"^\|[\s:|-]+\|$", lines[i + 1].strip()):
            head = split_row(s)
            i += 2
            body = []
            while i < n and lines[i].strip().startswith("|"):
                body.append(split_row(lines[i]))
                i += 1
            tbl = doc.add_table(rows=1, cols=len(head))
            tbl.style = "Table Grid"
            for c, txt in zip(tbl.rows[0].cells, head):
                c.text = ""
                add_runs(c.paragraphs[0], txt)
                for r in c.paragraphs[0].runs:
                    r.bold = True
            for row in body:
                for c, txt in zip(tbl.add_row().cells, row[:len(head)]):
                    c.text = ""
                    add_runs(c.paragraphs[0], txt)
            doc.add_paragraph()
            continue

        if s.startswith(">"):
            add_runs(doc.add_paragraph(style="Intense Quote"), s.lstrip("> ").strip())
            i += 1
            continue

        m = re.match(r"^(\d+)\.\s+(.*)", s)
        if m:
            i += 1
            rest, i = gather(lines, i, n)
            add_runs(doc.add_paragraph(style="List Number"),
                     (m.group(2) + " " + rest).strip())
            continue

        if re.match(r"^[-*+]\s+", s):
            head = re.sub(r"^[-*+]\s+", "", s)
            i += 1
            rest, i = gather(lines, i, n)
            add_runs(doc.add_paragraph(style="List Bullet"),
                     (head + " " + rest).strip())
            continue

        # Accumulate wrapped lines into ONE paragraph before inline parsing.
        # Markdown wraps bold spans across lines (**foo\nbar**); parsing line by
        # line leaves the opening ** unmatched and it renders as literal
        # asterisks in the output.
        buf = []
        while i < n:
            cur = lines[i].strip()
            if (not cur or cur.startswith(("```", "|", ">", "#", "!["))
                    or re.match(r"^[-*+]\s+|^\d+\.\s+", cur)
                    or (set(cur) <= {"-", "*", "_"} and len(cur) >= 3)):
                break
            buf.append(cur)
            i += 1
        add_runs(doc.add_paragraph(), " ".join(buf))

    doc.save(dst)
    print(f"wrote {dst}")


if __name__ == "__main__":
    main(sys.argv[1], sys.argv[2])
